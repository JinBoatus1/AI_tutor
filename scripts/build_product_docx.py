"""One-off script: generate PRODUCT.docx (run from repo root: python scripts/build_product_docx.py)."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out = root / "PRODUCT.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_heading("AI Tutor — Product Overview", 0)

    doc.add_heading("What It Is", level=1)
    doc.add_paragraph(
        "AI Tutor is a web-based learning companion built around a fixed curriculum (FOCS) and a "
        "textbook-grounded teaching style. It combines chat tutoring, structured learning flow, "
        "progress tracking, automated grading, and optional long-term memory per topic—so students "
        "get guided explanations that stay aligned with the book, not generic web answers."
    )

    doc.add_heading("Key Features & Why They Matter", level=1)

    doc.add_heading("1. Learning Mode (Interactive Tutor)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    for t in [
        "Chat interface: natural-language questions; optional images, screenshots, pasted figures, or PDF uploads (server renders first pages to images for multimodal understanding).",
        "Backend matches questions to FOCS topics, pulls relevant textbook pages from the official PDF, and shows navigable page images in a side panel.",
        "Short intake at session start (new vs review, progress, target sections) shapes the tutoring plan.",
        "Confidence score appears only for problem-solving turns—not for navigation or planning.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Generic chatbots often produce hallucinated or off-syllabus help. AI Tutor anchors answers to the real textbook and the course tree, improving trust and exam alignment."
    )

    doc.add_heading("2. Textbook & Topic Tree (FOCS)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    doc.add_paragraph(
        "FOCS.json defines chapters, sections, and page ranges; the FOCS PDF is the single source of truth. "
        "Topic matching and rendering use this tree so the left panel shows the correct chapter/section pages (full section range)."
    )
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Avoids wrong pages or misleading crops when students mention broad topics (e.g. “Induction”). The UI reflects actual book structure and page spans."
    )

    doc.add_heading("3. My Learning Bar (Progress Map)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    for t in [
        "Dedicated page showing the full FOCS tree.",
        "Learned vs not learned sections, driven by a per-student progress file.",
        "Click sections to toggle learned state; syncs with the same student ID as Learning Mode.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Progress is often invisible or scattered across chats. Here it is explicit, editable, and tied to the syllabus—useful for students and future adaptive logic."
    )

    doc.add_heading("4. Hidden Student Progress Bar (Backend)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    for t in [
        "Per-student JSON: current section, learned/planned sections, confusion counts.",
        "Heuristic updates from chat can expand to earlier chapters or earlier subsections in tree order.",
        "Injected into the tutor system prompt (not shown raw) with policies: stay in scope, bridge advanced ideas briefly, flag repeated confusion on “already learned” material.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Reduces teaching ahead of what the student has covered and supports gentle remediation when something marked learned still feels hard."
    )

    doc.add_heading("5. Auto Grader", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    doc.add_paragraph(
        "Students submit a grading prompt, free-text answer, and optional images; the backend returns model-generated feedback."
    )
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Scales draft feedback on written work when instructors or peers are unavailable—useful for practice; calibrate before any high-stakes use."
    )

    doc.add_heading("6. Session Memory (FOCS Subtopics)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    doc.add_paragraph(
        "For matched FOCS subtopics, summarized and full past Q&A can be stored; an optional tool lets the model load full history when summaries are insufficient."
    )
    doc.add_paragraph("Problem it solves", style="Heading 3")
    doc.add_paragraph(
        "Long threads forget earlier explanations. Scoped memory improves continuity within the same book section."
    )

    doc.add_heading("7. Curriculum Context (Legacy / Optional)", level=2)
    doc.add_paragraph("What it does", style="Heading 3")
    doc.add_paragraph(
        "A curriculum tree may live in local storage from older flows; Learning Mode can cross-match user wording for sidebar hints. The primary path is FOCS-native."
    )

    doc.add_heading("Problems Solved (Summary)", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Issue"
    hdr[2].text = "How AI Tutor Addresses It"
    rows = [
        ("Trust", "Generic AI ignores the real book", "PDF + FOCS tree + page images"),
        ("Structure", "Unclear position in the course", "My Learning Bar + hidden bar + intake"),
        ("Scope", "Tutor explains unseen material", "Bar-informed prompting + syllabus toggles"),
        ("Multimodal", "Questions in screenshots/PDFs", "Images + PDF→page images"),
        ("Feedback", "No quick written check", "Auto Grader with custom prompt"),
        ("Continuity", "New session, no context", "Subtopic memory + optional full history"),
    ]
    for area, issue, fix in rows:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = issue
        row[2].text = fix

    doc.add_heading("Technical Notes (High Level)", level=1)
    for t in [
        "Frontend: React + Vite; routes for Home, Learning Mode, Auto Grader, My Learning Bar.",
        "Backend: FastAPI; OpenAI-compatible chat; PyMuPDF for PDF text and rendering.",
        "Identity: anonymous student IDs in localStorage map to per-file progress bars on disk.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Positioning One-Liner", level=1)
    doc.add_paragraph(
        "AI Tutor is a textbook-first, syllabus-aware AI study companion: it answers in the language of your FOCS course, "
        "shows the real pages, remembers your progress, and keeps explanations inside what you’ve actually learned—while still "
        "supporting images, PDFs, grading help, and session memory on demand."
    )

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
